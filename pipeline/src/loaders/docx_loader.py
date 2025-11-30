"""Loader for DOCX files."""

from __future__ import annotations

from pathlib import Path
from typing import List

from rag.models import Document

from .base import DocumentLoader, register_loader

try:
    import docx  # type: ignore
except Exception:  # pragma: no cover - optional dependency
    docx = None


@register_loader([".docx"])
class DocxLoader(DocumentLoader):
    """Load text from .docx files.

    Uses `python-docx` if installed; otherwise raises a clear error so the
    user knows to install the extra dependency.
    """

    def load(self, path: Path) -> List[Document]:
        if docx is None:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                "Install it with `pip install python-docx`."
            )

        doc_obj = docx.Document(str(path))
        paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        doc = Document(
            id=str(path),
            path=path,
            text=text,
            metadata={"source": "docx", "suffix": ".docx"},
        )
        return [doc]


